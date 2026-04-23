"""
generators/docx_generator.py — Build per-file and master DOCX output documents.

Uses python-docx to produce well-formatted Word documents with:
  - Cover section per file with metadata
  - Diagram explanations formatted with clear headings and step structure
  - Prose sections with preserved hierarchy
  - A master aggregated document combining all results
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from  app.processors.diagram_explainer import ExplainedContent
from  app.utils.logger import get_logger

logger = get_logger(__name__)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _add_horizontal_rule(doc: Document) -> None:
    """Add a paragraph with a bottom border acting as a horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _style_heading(para, level: int = 1) -> None:
    """Apply explicit font styling to a heading paragraph."""
    run = para.runs[0] if para.runs else para.add_run()
    run.font.bold = True
    run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


def _add_file_header(doc: Document, file_name: str, page_count: int) -> None:
    """Add a styled file-level section header."""
    _add_horizontal_rule(doc)
    h = doc.add_heading(f"📄  {file_name}", level=1)
    _style_heading(h, 1)
    meta = doc.add_paragraph(f"Source file: {file_name}  |  Sections processed: {page_count}")
    meta.runs[0].font.italic = True
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    doc.add_paragraph()


def _add_diagram_section(doc: Document, content: ExplainedContent) -> None:
    """Render one diagram explanation into the document."""
    # Section heading
    h2 = doc.add_heading(
        f"🔷  {content.section_title}  [{content.diagram_type.replace('_', ' ').title()}]",
        level=2,
    )
    _style_heading(h2, 2)

    # Confidence badge
    conf_para = doc.add_paragraph()
    conf_run = conf_para.add_run(
        f"Detection confidence: {content.confidence:.0%}  |  Type: {content.diagram_type}"
    )
    conf_run.font.italic = True
    conf_run.font.size = Pt(9)
    conf_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Explanation body
    _render_explanation_body(doc, content.explanation)
    doc.add_paragraph()


def _add_prose_section(doc: Document, content: ExplainedContent) -> None:
    """Render a non-diagram page summary into the document."""
    h2 = doc.add_heading(f"📝  {content.section_title}", level=2)
    _style_heading(h2, 2)
    _render_explanation_body(doc, content.explanation)
    doc.add_paragraph()


def _render_explanation_body(doc: Document, text: str) -> None:
    """
    Intelligently render LLM output text into the document:
      - Lines starting with digits + . or ) → numbered list style
      - Lines starting with - or • → bullet style
      - Lines starting with # → sub-heading
      - Everything else → normal paragraph
    """
    num_pattern = re.compile(r"^\s*(\d+[a-z]?[\.\)])\s+")
    bullet_pattern = re.compile(r"^\s*[-•*]\s+")
    heading_pattern = re.compile(r"^\s*#{1,3}\s+(.+)")

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        h_match = heading_pattern.match(stripped)
        if h_match:
            p = doc.add_paragraph(h_match.group(1))
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            continue

        if num_pattern.match(stripped):
            p = doc.add_paragraph(stripped, style="List Number")
            p.paragraph_format.left_indent = Inches(0.25)
            continue

        if bullet_pattern.match(stripped):
            clean = bullet_pattern.sub("", stripped)
            p = doc.add_paragraph(clean, style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            continue

        # Normal paragraph — detect inline bold (**text**)
        p = doc.add_paragraph()
        _add_inline_formatted_run(p, stripped)


def _add_inline_formatted_run(para, text: str) -> None:
    """
    Split text on **bold** markers and add formatted runs.
    Falls back to a single plain run if no markers found.
    """
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.size = Pt(11)
        if i % 2 == 1:  # Odd indices are inside ** **
            run.bold = True


# ── Per-file document generator ───────────────────────────────────────────────

def generate_file_docx(
    file_name: str,
    contents: List[ExplainedContent],
    output_path: Path,
) -> Path:
    """
    Generate a DOCX for a single source file.

    Parameters
    ----------
    file_name : str
        Original file name (used in headings).
    contents : List[ExplainedContent]
        Ordered list of explained pages/sections for this file.
    output_path : Path
        Where to save the generated .docx file.

    Returns
    -------
    Path
        The saved document path.
    """
    doc = Document()

    # ── Document-level styling ────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover header ──────────────────────────────────────────────────────────
    _add_file_header(doc, file_name, len(contents))

    diagram_count = sum(1 for c in contents if c.content_type == "diagram")
    prose_count = sum(1 for c in contents if c.content_type == "prose")

    summary_para = doc.add_paragraph(
        f"Summary: {diagram_count} diagram section(s) and {prose_count} prose section(s) extracted."
    )
    summary_para.runs[0].font.bold = True

    doc.add_paragraph()

    # ── Content sections ──────────────────────────────────────────────────────
    for content in contents:
        if content.content_type == "diagram":
            _add_diagram_section(doc, content)
        elif content.content_type == "prose":
            _add_prose_section(doc, content)
        else:  # error
            h = doc.add_heading(f"⚠️  {content.section_title} — Processing Error", level=2)
            doc.add_paragraph(content.explanation)

    doc.save(str(output_path))
    logger.info("File DOCX saved: %s (%d sections)", output_path.name, len(contents))
    return output_path


# ── Master aggregated document ────────────────────────────────────────────────

def generate_master_docx(
    results: List[dict],   # [{"file_name": str, "contents": List[ExplainedContent]}]
    output_path: Path,
) -> Path:
    """
    Generate the master DOCX combining outputs from all processed files.

    Parameters
    ----------
    results : List[dict]
        Each entry has keys "file_name" and "contents".
    output_path : Path
        Where to save the master .docx.

    Returns
    -------
    Path
        The saved master document path.
    """
    doc = Document()

    # ── Document-level styling ────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover page ────────────────────────────────────────────────────────────
    title_para = doc.add_heading("Document Analysis Master Report", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph(
        f"Generated by Diagram Processor  |  Files processed: {len(results)}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # ── Table of contents placeholder ────────────────────────────────────────
    doc.add_heading("Files Processed", level=1)
    for i, result in enumerate(results, 1):
        file_name = result["file_name"]
        n_sections = len(result["contents"])
        n_diagrams = sum(1 for c in result["contents"] if c.content_type == "diagram")
        p = doc.add_paragraph(
            f"{i}. {file_name}  —  {n_sections} section(s), {n_diagrams} diagram(s)"
        )
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # ── Per-file content ──────────────────────────────────────────────────────
    for result in results:
        file_name = result["file_name"]
        contents: List[ExplainedContent] = result["contents"]

        _add_file_header(doc, file_name, len(contents))

        for content in contents:
            if content.content_type == "diagram":
                _add_diagram_section(doc, content)
            elif content.content_type == "prose":
                _add_prose_section(doc, content)
            else:
                doc.add_heading(f"⚠️  {content.section_title} — Error", level=2)
                doc.add_paragraph(content.explanation)

        doc.add_page_break()

    doc.save(str(output_path))
    logger.info(
        "Master DOCX saved: %s (%d file(s))", output_path.name, len(results)
    )
    return output_path