"""
parsers/docx_parser.py — Extract text and embedded images from Word documents.

Paragraphs are read in document order.  Inline images (drawings / shapes)
are extracted from the document XML relationships.  Heading styles are used
to determine section titles.  .doc files are converted via LibreOffice.
"""

from __future__ import annotations

import io
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional
from xml.etree import ElementTree as ET

import docx
from docx import Document
from docx.oxml.ns import qn

from app.parsers.base_parser import BaseParser, ImageBlock, PageContent, ParserError
from app.config import settings
from app.utils.logger import get_logger

logger = get_logger(__name__)

_HEADING_STYLES = {
    "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "Title", "Subtitle", "heading 1", "heading 2",
}

_DIAGRAM_KEYWORDS = {
    "start", "end", "yes", "no", "decision", "process", "flow",
    "step", "if", "then", "else", "loop", "condition", "branch",
    "gateway", "→", "->", "⇒",
}


def _convert_doc_to_docx(doc_path: Path) -> Optional[Path]:
    """Convert legacy .doc → .docx using LibreOffice."""
    try:
        tmp_dir = Path(tempfile.mkdtemp())
        result = subprocess.run(
            [
                "libreoffice", "--headless",
                "--convert-to", "docx",
                "--outdir", str(tmp_dir),
                str(doc_path),
            ],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0:
            converted = list(tmp_dir.glob("*.docx"))
            if converted:
                logger.info("Converted %s → %s", doc_path.name, converted[0].name)
                return converted[0]
        logger.warning("LibreOffice doc→docx failed: %s", result.stderr)
    except FileNotFoundError:
        logger.warning("LibreOffice not found; cannot convert .doc files")
    except Exception as exc:
        logger.warning("DOC conversion error: %s", exc)
    return None


def _extract_images_from_docx(doc_obj: Document) -> List[ImageBlock]:
    """Extract all inline images from a python-docx Document object."""
    images: List[ImageBlock] = []
    rels = doc_obj.part.rels
    for rel_id, rel in rels.items():
        if "image" in rel.reltype:
            try:
                blob = rel.target_part.blob
                content_type = rel.target_part.content_type  # e.g. "image/png"
                images.append(
                    ImageBlock(
                        image_bytes=blob,
                        mime_type=content_type,
                        page_number=0,  # DOCX has no page concept without rendering
                    )
                )
            except Exception as exc:
                logger.warning("Could not extract image %s: %s", rel_id, exc)
    return images


class DocxParser(BaseParser):
    """Parse .docx (and .doc via LibreOffice conversion) files."""

    def parse(self) -> List[PageContent]:
        target = self.file_path

        if self.file_path.suffix.lower() == ".doc":
            converted = _convert_doc_to_docx(self.file_path)
            if converted is None:
                raise ParserError(
                    f"Cannot parse legacy .doc without LibreOffice: {self.file_name}"
                )
            target = converted

        try:
            doc_obj = Document(str(target))
        except Exception as exc:
            raise ParserError(f"Cannot open document {self.file_name}: {exc}") from exc

        # Extract images once at the document level
        all_images = _extract_images_from_docx(doc_obj)

        # Group paragraphs into logical sections (split at H1 boundaries)
        sections: List[PageContent] = []
        current_title = self.file_name
        current_texts: List[str] = []
        pseudo_page = 1

        def _flush_section():
            nonlocal pseudo_page
            if current_texts:
                text = "\n".join(current_texts)
                kw_hits = sum(1 for kw in _DIAGRAM_KEYWORDS if kw in text.lower())
                sections.append(
                    PageContent(
                        page_number=pseudo_page,
                        text=text,
                        images=[],  # images assigned below
                        section_title=current_title,
                        is_diagram_page=kw_hits >= settings.diagram_keyword_threshold,
                        raw_metadata={"source": "docx"},
                    )
                )
                pseudo_page += 1

        for para in doc_obj.paragraphs:
            style_name = para.style.name if para.style else ""
            para_text = para.text.strip()

            if style_name in _HEADING_STYLES and para_text:
                _flush_section()
                current_title = para_text[:120]
                current_texts = [para_text]
            else:
                if para_text:
                    current_texts.append(para_text)

        # Also grab text from tables
        for table in doc_obj.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    current_texts.append(" | ".join(row_texts))

        _flush_section()

        # Distribute images across sections proportionally
        if sections and all_images:
            per_section = max(1, len(all_images) // len(sections))
            for i, section in enumerate(sections):
                start = i * per_section
                end = start + per_section if i < len(sections) - 1 else len(all_images)
                section.images = all_images[start:end]

        # Ensure at least one PageContent even for empty docs
        if not sections:
            sections.append(
                PageContent(
                    page_number=1,
                    text="[Document appears to be empty]",
                    section_title=self.file_name,
                    raw_metadata={"source": "docx"},
                )
            )

        logger.info("DOCX parsed: %s — %d section(s)", self.file_name, len(sections))
        return sections